import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    # Custom styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []

    for line in lines:
        line = line.strip()
        
        # 1. Handle Headings
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            continue
        
        # 2. Handle Tables
        if '|' in line:
            # Skip separator lines like | :--- | :--- |
            if '---' in line:
                continue
            
            in_table = True
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                table_data.append(cells)
            continue
        else:
            if in_table:
                # End of table, process it
                if table_data:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Table Grid'
                    for r, row_data in enumerate(table_data):
                        for c, cell_text in enumerate(row_data):
                            # Handle bold in cells
                            clean_text = cell_text.replace('**', '').replace('__', '')
                            table.cell(r, c).text = clean_text
                
                in_table = False
                table_data = []
            
        # 3. Handle Horizontal Rules
        if line == '---':
            doc.add_paragraph("_" * 60)
            continue
            
        # 4. Handle Alerts (Simplified)
        if line.startswith('> [!'):
            p = doc.add_paragraph()
            run = p.add_run(line.split(']')[0][4:].replace('!', '') + ": ")
            run.bold = True
            run = p.add_run(line.split(']')[1].strip())
            run.italic = True
            continue
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.paragraph_format.left_indent = Pt(18)
            continue

        # 5. Handle Bullet Points
        if line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
            continue
        
        # 6. Regular Paragraphs (preserving bold)
        if line:
            p = doc.add_paragraph()
            # Extremely simple bold parsing
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True

    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

if __name__ == "__main__":
    convert_md_to_docx('Project_Statistical_Report.md', 'Project_Statistical_Report.docx')
